import os
import tkinter as tk
from tkinter import filedialog, simpledialog, messagebox
from docx import Document


# --------------------------------------------------
# ניקוי אגרסיבי של טקסט ל־XML (python-docx safe)
# --------------------------------------------------
def clean_xml_text(text: str) -> str:
    return "".join(
        c for c in text
        if ord(c) in (0x9, 0xA, 0xD)
        or 0x20 <= ord(c) <= 0xD7FF
        or 0xE000 <= ord(c) <= 0xFFFD
        or 0x10000 <= ord(c) <= 0x10FFFF
    )


def add_safe_paragraph(document, text: str, chunk_size: int = 1000):
    paragraph = document.add_paragraph()
    for i in range(0, len(text), chunk_size):
        paragraph.add_run(text[i:i + chunk_size])


# --------------------------------------------------
# GUI helpers
# --------------------------------------------------
def browse_folder():
    folder = filedialog.askdirectory()
    if folder:
        folder_path.set(folder)


SKIP_DIRS = {
    "node_modules", ".git", "__pycache__", "venv", ".venv",
    "dist", "build", ".next", ".nuxt", "out", "coverage",
    ".mypy_cache", ".pytest_cache", ".tox", "vendor",
}


def collect_files(folder, extensions, file_filter=None):
    """
    file_filter: פונקציה שמקבלת file_path ומחזירה True / False
    """
    document = Document()
    document.add_heading("Collected Source Files", level=1)
    file_count = 0

    for root, dirs, files in os.walk(folder):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS]

        for file in files:
            if os.path.splitext(file)[1].lower() not in extensions:
                continue

            file_path = os.path.join(root, file)

            if file_filter and not file_filter(file_path):
                continue

            file_count += 1
            document.add_heading(file_path, level=2)

            try:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    raw_content = f.read()
            except Exception as e:
                raw_content = f"[שגיאה בקריאת הקובץ: {e}]"

            add_safe_paragraph(document, clean_xml_text(raw_content))

    return document, file_count


# --------------------------------------------------
# סריקה כללית לפי סיומות
# --------------------------------------------------
def start_scan():
    folder = folder_path.get()
    if not folder:
        messagebox.showerror("שגיאה", "לא נבחרה תיקייה")
        return

    ext_input = simpledialog.askstring(
        "סיומות קבצים",
        "הכנס סיומות מופרדות בפסיק (לדוגמה: js, jsx, ts, py):"
    )
    if not ext_input:
        return

    extensions = {
        f".{ext.strip().lower().lstrip('.')}"
        for ext in ext_input.split(",")
        if ext.strip()
    }

    document, file_count = collect_files(folder, extensions)

    if file_count == 0:
        messagebox.showinfo("תוצאה", "לא נמצאו קבצים מתאימים")
        return

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word file", "*.docx")]
    )
    if save_path:
        document.save(save_path)
        messagebox.showinfo("הצלחה", f"הועתקו {file_count} קבצים בהצלחה")


# --------------------------------------------------
# פילטר: קבצי Python עם class
# --------------------------------------------------
def has_class(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return any(
                line.lstrip().startswith("class ")
                for line in f
            )
    except Exception:
        return False


def start_scan_py():
    folder = folder_path.get()
    if not folder:
        messagebox.showerror("שגיאה", "לא נבחרה תיקייה")
        return

    document, file_count = collect_files(
        folder,
        extensions={".py"},
        file_filter=has_class
    )

    if file_count == 0:
        messagebox.showinfo("תוצאה", "לא נמצאו קבצי Python עם מחלקות")
        return

    save_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word file", "*.docx")]
    )
    if save_path:
        document.save(save_path)
        messagebox.showinfo(
            "הצלחה",
            f"הועתקו {file_count} קבצי Python עם מחלקות"
        )


# --------------------------------------------------
# GUI setup
# --------------------------------------------------
root = tk.Tk()
root.title("Source Code → Word Collector")

folder_path = tk.StringVar()

frame = tk.Frame(root, padx=20, pady=20)
frame.pack()

tk.Label(frame, text="בחר תיקייה לסריקה:").pack(anchor="w")

tk.Entry(frame, textvariable=folder_path, width=60).pack(side="left")
tk.Button(frame, text="Browse", command=browse_folder).pack(side="left", padx=5)

tk.Button(
    root,
    text="Scan & Export to Word (לפי סיומות)",
    command=start_scan,
    bg="#cce7ff"
).pack(pady=(20, 5))

tk.Button(
    root,
    text="Export Python files with classes only",
    command=start_scan_py,
    bg="#d4edda"
).pack(pady=(0, 20))

root.mainloop()